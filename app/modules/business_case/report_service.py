"""A business case as a document — PDF, Word, Excel and a shareable page.

The business case is the one business-architecture artefact that already *is* a
document: problem, options, recommendation, costs, benefits, risks. Until now it
existed only as a web form, which meant the version that reached an investment
board was whatever somebody retyped into their own template.

One builder, four renderers. ``build()`` gathers the case, its three optional
links and the financial cross-check exactly once, so the PDF, the Word file and
the workbook cannot present different capex for the same case in the same
meeting.

Two rules this file exists to hold:

**It never invents a number.** Every financial column on ``BusinessCase`` is
nullable. A case with no capex entered shows an em dash and an empty workbook
cell, never 0 — a zero capex is a claim that the change is free, and somebody
approves against it. Narrative sections that have not been written say so in
words rather than appearing as an empty heading that reads as "nothing to say".

**It never writes.** ``aggregate_financials`` is called with
``apply_missing=False``. Exporting a document must not mutate the document: the
linked capability / initiative / solution figures appear as a separate
cross-check section, clearly labelled as coming from elsewhere, and the case's
own fields stay exactly as the author left them.
"""

from __future__ import annotations

import io
from datetime import datetime
from decimal import Decimal
from typing import Any, Dict, List, Optional

EM_DASH = "—"

# The document sections, in the order a reader expects them.
NARRATIVE_SECTIONS = (
    ("problem_statement", "Problem statement"),
    ("options_considered", "Options considered"),
    ("recommended_option", "Recommendation"),
    ("expected_benefits", "Expected benefits"),
    ("key_risks", "Key risks"),
)

# label, attribute, unit. Unit is a suffix, not a currency symbol: the model
# stores a bare amount and no currency, and stamping one on would be invention.
FINANCIAL_FIELDS = (
    ("Capital expenditure", "capex", None),
    ("Operating expenditure (annual)", "opex_annual", None),
    ("Total cost of ownership (3 year)", "tco_3yr", None),
    ("Financial benefit (annual)", "financial_benefit_annual", None),
    ("Return on investment", "roi_percentage", "%"),
    ("Payback period", "payback_months", "months"),
)


class BusinessCaseReportError(Exception):
    """A renderer could not produce its format on this deployment."""


class BusinessCaseReportService:
    """Build and render the business-case report."""

    # ------------------------------------------------------------------
    # Gathering
    # ------------------------------------------------------------------

    @staticmethod
    def build(
        business_case_id: int, organisation_name: Optional[str] = None
    ) -> Optional[Dict[str, Any]]:
        """Collect everything the report needs, once. None if no such case.

        Runs inside the request's tenant context: BusinessCase is a TenantMixin
        model, so the organisation predicate is applied by do_orm_execute and
        must not be written here as well.

        ``filter(...).first()`` rather than ``.get()`` deliberately —
        ``Query.get()`` is scoped only on an identity-map miss, and on a hit it
        returns the cached row without emitting SQL, so no tenant filter runs
        (see CLAUDE.md).
        """
        from app.models.business_case import BusinessCase

        from . import service

        case = BusinessCase.query.filter(BusinessCase.id == business_case_id).first()
        if case is None:
            return None

        # apply_missing=False: rendering a document must not write to it.
        try:
            aggregation = service.aggregate_financials(case, apply_missing=False)
        except Exception:  # noqa: BLE001 - the case is the source of truth
            # The cross-check is supplementary. Losing it must not lose the
            # document, but it must not be replaced with invented figures
            # either — the section simply reports that it is unavailable.
            aggregation = None

        narrative = [
            {"key": key, "label": label, "text": getattr(case, key) or None}
            for key, label in NARRATIVE_SECTIONS
        ]

        financials = [
            {
                "label": label,
                "key": key,
                "value": _number(getattr(case, key)),
                "unit": unit,
            }
            for label, key, unit in FINANCIAL_FIELDS
        ]

        links = [
            {
                "label": "Capability",
                "name": case.capability.name if case.capability else None,
            },
            {
                "label": "Strategic initiative",
                "name": (
                    case.strategic_initiative.name if case.strategic_initiative else None
                ),
            },
            {
                "label": "Solution",
                "name": case.solution.name if case.solution else None,
            },
        ]

        cross_checks: List[Dict[str, Any]] = []
        sources: List[Dict[str, Any]] = []
        if aggregation:
            label_by_key = {key: label for label, key, _unit in FINANCIAL_FIELDS}
            for key, delta in sorted((aggregation.get("cross_checks") or {}).items()):
                cross_checks.append(
                    {
                        "label": label_by_key.get(key, key),
                        "entered": _number(delta.get("current")),
                        "linked": _number(delta.get("suggested")),
                        "difference": _difference(
                            delta.get("current"), delta.get("suggested")
                        ),
                    }
                )
            for source_key, heading in (
                ("strategic_initiative", "Strategic initiative"),
                ("capability_cost_allocation", "Capability cost allocation"),
                ("unified_capability", "Capability (unified model)"),
                ("solution", "Solution"),
            ):
                payload = aggregation.get(source_key)
                if payload:
                    # "fields", not "values": Jinja resolves `source.values` to
                    # dict.values, the bound method, before it ever looks for a
                    # key of that name.
                    sources.append({"heading": heading, "fields": payload})

        return {
            "organisation": organisation_name,
            "generated_at": datetime.utcnow(),
            "case": {
                "id": case.id,
                "title": case.title,
                "description": case.description,
                "status": case.status,
                "author": case.created_by.full_name() if case.created_by else None,
                "created_at": case.created_at,
                "updated_at": case.updated_at,
            },
            "links": links,
            "narrative": narrative,
            "financials": financials,
            "cross_checks": cross_checks,
            "sources": sources,
            "aggregation_available": aggregation is not None,
            "totals": {
                "sections_written": sum(1 for row in narrative if row["text"]),
                "sections": len(narrative),
                "financials_entered": sum(
                    1 for row in financials if row["value"] is not None
                ),
                "financials": len(financials),
                "links": sum(1 for row in links if row["name"]),
            },
        }

    # ------------------------------------------------------------------
    # Rendering
    # ------------------------------------------------------------------

    @staticmethod
    def render_pdf(report: Dict[str, Any]) -> bytes:
        """HTML → PDF, following the pattern already used for the SAD export."""
        from flask import render_template

        html = render_template("exports/business_case_report.html", **report, for_pdf=True)

        try:
            from weasyprint import HTML

            return HTML(string=html).write_pdf()
        except (ImportError, OSError):
            # OSError as well as ImportError: WeasyPrint is a CFFI binding over
            # native GTK/Pango, and when the wheel is installed without those
            # libraries — normal on Windows and on slim images — the import
            # raises OSError, not ImportError.
            try:
                import pdfkit

                return pdfkit.from_string(
                    html,
                    False,
                    options={
                        "page-size": "A4",
                        "margin-top": "18mm",
                        "margin-right": "16mm",
                        "margin-bottom": "18mm",
                        "margin-left": "16mm",
                        "encoding": "UTF-8",
                        "enable-local-file-access": None,
                    },
                )
            except Exception as exc:  # noqa: BLE001
                raise BusinessCaseReportError(
                    "PDF rendering is unavailable on this deployment (neither "
                    "WeasyPrint's native libraries nor wkhtmltopdf are present). "
                    "Word and Excel export still work."
                ) from exc

    @staticmethod
    def render_html(report: Dict[str, Any]) -> str:
        """The same document as a page — a shareable link, and what the PDF renders."""
        from flask import render_template

        return render_template(
            "exports/business_case_report.html", **report, for_pdf=False
        )

    @classmethod
    def render_docx(cls, report: Dict[str, Any]) -> bytes:
        """Word, because an investment board edits before it approves."""
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - python-docx is pinned
            raise BusinessCaseReportError(
                "Word export is unavailable on this deployment (python-docx is missing)."
            ) from exc

        case = report["case"]
        document = Document()

        document.add_heading(case["title"] or "Business Case", level=0)
        subtitle = document.add_paragraph()
        subtitle.add_run(report.get("organisation") or "Business architecture").bold = True
        subtitle.add_run(
            f"  ·  Business case  ·  {(case['status'] or 'draft').capitalize()}"
            f"  ·  {report['generated_at'].strftime('%d %B %Y')}"
        )
        if case["description"]:
            document.add_paragraph(case["description"])

        document.add_heading("At a glance", level=1)
        glance = document.add_table(rows=0, cols=2)
        glance.style = "Light Grid Accent 1"
        for label, value in (
            ("Status", (case["status"] or "draft").capitalize()),
            ("Author", case["author"] or EM_DASH),
            ("Created", _date(case["created_at"])),
            ("Last updated", _date(case["updated_at"])),
        ):
            cells = glance.add_row().cells
            cells[0].text = label
            cells[1].text = value
        for link in report["links"]:
            cells = glance.add_row().cells
            cells[0].text = link["label"]
            cells[1].text = link["name"] or "Not linked"

        for section in report["narrative"]:
            document.add_heading(section["label"], level=1)
            if section["text"]:
                for paragraph in str(section["text"]).splitlines():
                    if paragraph.strip():
                        document.add_paragraph(paragraph)
            else:
                empty = document.add_paragraph()
                empty.add_run(
                    "Not documented. This section has not been written yet."
                ).italic = True

        document.add_heading("Financial summary", level=1)
        financials = document.add_table(rows=1, cols=2)
        financials.style = "Light Grid Accent 1"
        financials.rows[0].cells[0].text = "Measure"
        financials.rows[0].cells[1].text = "Value"
        for row in report["financials"]:
            cells = financials.add_row().cells
            cells[0].text = row["label"]
            cells[1].text = cls._fmt(row["value"], row["unit"])

        note = document.add_paragraph()
        note.add_run(
            "An em dash means the figure has not been entered. It is not a zero: "
            "a zero cost is a claim that the change is free."
        ).italic = True

        document.add_heading("Cross-check against linked records", level=1)
        if not report["aggregation_available"]:
            unavailable = document.add_paragraph()
            unavailable.add_run(
                "The linked-record cross-check could not be computed for this "
                "case, so no comparison figures are shown."
            ).italic = True
        elif report["cross_checks"]:
            comparison = document.add_table(rows=1, cols=4)
            comparison.style = "Light Grid Accent 1"
            for idx, header in enumerate(
                ("Measure", "In this case", "From linked records", "Difference")
            ):
                comparison.rows[0].cells[idx].text = header
            for row in report["cross_checks"]:
                cells = comparison.add_row().cells
                cells[0].text = row["label"]
                cells[1].text = cls._fmt(row["entered"], None)
                cells[2].text = cls._fmt(row["linked"], None)
                cells[3].text = cls._fmt(row["difference"], None)
        else:
            clean = document.add_paragraph()
            clean.add_run(
                "No figure in this case diverges from the linked capability, "
                "initiative or solution records."
            ).italic = True

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @classmethod
    def render_xlsx(cls, report: Dict[str, Any]) -> bytes:
        """Excel, because the numbers get modelled before they get presented."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Font
        except ImportError as exc:  # pragma: no cover - openpyxl is pinned
            raise BusinessCaseReportError(
                "Excel export is unavailable on this deployment (openpyxl is missing)."
            ) from exc

        case = report["case"]
        workbook = Workbook()

        summary = workbook.active
        summary.title = "Summary"
        summary.append([case["title"] or "Business Case"])
        summary["A1"].font = Font(bold=True, size=14)
        summary.append([])
        for label, value in (
            ("Generated", report["generated_at"].strftime("%Y-%m-%d %H:%M UTC")),
            ("Organisation", report.get("organisation") or ""),
            ("Status", case["status"] or ""),
            ("Author", case["author"] or ""),
            ("Created", _date(case["created_at"], blank=True)),
            ("Last updated", _date(case["updated_at"], blank=True)),
            ("Description", case["description"] or ""),
        ):
            summary.append([label, value])
        for link in report["links"]:
            summary.append([link["label"], link["name"] or ""])
        _widths(summary, [30, 60])

        financials = workbook.create_sheet("Financials")
        financials.append(["Measure", "Value", "Unit"])
        for cell in financials[1]:
            cell.font = Font(bold=True)
        for row in report["financials"]:
            financials.append(
                [
                    row["label"],
                    # Empty, not zero: a 0 in a cost column reads as a measured
                    # figure and gets summed into a portfolio total.
                    "" if row["value"] is None else row["value"],
                    row["unit"] or "",
                ]
            )
        _widths(financials, [36, 18, 12])

        narrative = workbook.create_sheet("Narrative")
        narrative.append(["Section", "Content"])
        for cell in narrative[1]:
            cell.font = Font(bold=True)
        for section in report["narrative"]:
            narrative.append([section["label"], section["text"] or ""])
        _widths(narrative, [28, 110])
        for row in narrative.iter_rows(min_row=2, min_col=2, max_col=2):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        cross = workbook.create_sheet("Cross-check")
        cross.append(["Measure", "In this case", "From linked records", "Difference"])
        for cell in cross[1]:
            cell.font = Font(bold=True)
        for row in report["cross_checks"]:
            cross.append(
                [
                    row["label"],
                    "" if row["entered"] is None else row["entered"],
                    "" if row["linked"] is None else row["linked"],
                    "" if row["difference"] is None else row["difference"],
                ]
            )
        _widths(cross, [36, 18, 22, 16])

        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _fmt(value: Optional[float], unit: Optional[str]) -> str:
        """A figure nobody entered is an em dash. It is never 0."""
        if value is None:
            return EM_DASH
        if float(value).is_integer():
            rendered = f"{int(value):,}"
        else:
            rendered = f"{value:,.2f}"
        return f"{rendered} {unit}" if unit else rendered


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _number(value) -> Optional[float]:
    """Decimal/int/float -> float, and anything absent stays absent."""
    if value is None:
        return None
    if isinstance(value, Decimal):
        return float(value)
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _difference(current, suggested) -> Optional[float]:
    left, right = _number(current), _number(suggested)
    if left is None or right is None:
        return None
    return left - right


def _date(value, blank: bool = False) -> str:
    if value is None:
        return "" if blank else EM_DASH
    return value.strftime("%d %B %Y")


def _widths(sheet, widths) -> None:
    for idx, width in enumerate(widths, start=1):
        sheet.column_dimensions[sheet.cell(row=1, column=idx).column_letter].width = width
