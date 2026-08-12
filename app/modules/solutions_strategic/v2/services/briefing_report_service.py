"""The EA briefing as a document — PDF, Word, Excel and a shareable page.

``EnterpriseBriefingService`` already computes the findings and persists them;
what it could not do was hand somebody the result. A briefing whose only form is
a web page is a briefing that gets screenshotted into a slide, which is where
the evidence link — the whole point of the artefact — is lost.

One builder, four renderers. ``build()`` reads the persisted briefing once and
each renderer formats that same structure, so the counts on the cover of the
PDF and the counts in the workbook cannot disagree.

Nothing here recomputes and nothing here invents. The findings are read back
exactly as they were written at generation time, including the ``action_url``
that says where each one can be verified; a finding missing a field renders an
em dash rather than a plausible substitute. Regenerating is an explicit user
action on the briefing page — exporting a document must not silently produce a
different set of findings from the one on screen.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

EM_DASH = "—"

SEVERITY_ORDER = {"critical": 0, "high": 1, "info": 2}


class BriefingReportError(Exception):
    """A renderer could not produce its format on this deployment."""


class BriefingReportService:
    """Build and render the Enterprise Architecture briefing report."""

    # ------------------------------------------------------------------
    # Gathering
    # ------------------------------------------------------------------

    @staticmethod
    def build(
        briefing_id: Optional[int] = None,
        organisation_name: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:
        """Read one persisted briefing — the newest by default. None if there is none.

        Note for anyone extending this: ``EnterpriseBriefing`` does **not**
        inherit ``TenantMixin``, so unlike the capability, value-stream and
        business-case reports there is no organisation predicate injected by
        do_orm_execute here. That is a property of the existing model, not
        something this exporter introduces or can fix locally — the same rows
        are already visible on /solutions/briefings.
        """
        from app.models.strategic import EnterpriseBriefing

        from .enterprise_briefing_service import EnterpriseBriefingService

        if briefing_id is None:
            briefing = EnterpriseBriefingService.latest()
        else:
            briefing = EnterpriseBriefing.query.filter(
                EnterpriseBriefing.id == briefing_id
            ).first()
        if briefing is None:
            return None

        findings: List[Dict[str, Any]] = []
        for raw in briefing.findings or []:
            if not isinstance(raw, dict):
                continue
            findings.append(
                {
                    "category": raw.get("category"),
                    "severity": raw.get("severity"),
                    "title": raw.get("title"),
                    "detail": raw.get("detail"),
                    "evidence": raw.get("evidence"),
                    "action_label": raw.get("action_label"),
                    "action_url": raw.get("action_url"),
                }
            )
        findings.sort(
            key=lambda f: (
                SEVERITY_ORDER.get(f["severity"], 3),
                (f["category"] or ""),
                (f["title"] or ""),
            )
        )

        by_category: Dict[str, Dict[str, Any]] = {}
        by_severity: Dict[str, int] = {}
        for finding in findings:
            category = finding["category"] or "uncategorised"
            bucket = by_category.setdefault(
                category, {"name": category, "count": 0, "flagged": 0}
            )
            bucket["count"] += 1
            if finding["severity"] in ("critical", "high"):
                bucket["flagged"] += 1
            severity = finding["severity"] or "unclassified"
            by_severity[severity] = by_severity.get(severity, 0) + 1

        history = [
            {
                "id": row.id,
                "generated_at": row.generated_at,
                "source": row.source,
                "headline": row.headline,
                "finding_count": row.finding_count,
                "flagged_count": row.flagged_count,
            }
            for row in EnterpriseBriefingService.history()
        ]

        return {
            "organisation": organisation_name,
            "generated_at": datetime.utcnow(),
            "briefing": {
                "id": briefing.id,
                "headline": briefing.headline,
                "summary": briefing.summary,
                "source": briefing.source,
                "briefed_at": briefing.generated_at,
                # These are the counts stored with the briefing. They are not
                # recomputed from the findings list: the stored number is what
                # the platform published at the time and is what the page shows.
                "finding_count": briefing.finding_count,
                "flagged_count": briefing.flagged_count,
            },
            "findings": findings,
            "categories": sorted(
                by_category.values(), key=lambda c: (-c["flagged"], -c["count"], c["name"])
            ),
            "severities": [
                {"name": name, "count": by_severity[name]}
                for name in sorted(
                    by_severity, key=lambda s: SEVERITY_ORDER.get(s, 3)
                )
            ],
            "history": history,
            "totals": {
                "findings_listed": len(findings),
                "categories": len(by_category),
            },
        }

    # ------------------------------------------------------------------
    # Rendering
    # ------------------------------------------------------------------

    @staticmethod
    def render_pdf(report: Dict[str, Any]) -> bytes:
        """HTML → PDF, following the pattern already used for the SAD export."""
        from flask import render_template

        html = render_template("exports/ea_briefing_report.html", **report, for_pdf=True)

        try:
            from weasyprint import HTML

            return HTML(string=html).write_pdf()
        except (ImportError, OSError):
            # OSError as well as ImportError: WeasyPrint is a CFFI binding over
            # native GTK/Pango, and when the wheel is installed without those
            # libraries — normal on Windows and on slim images — the import
            # raises OSError, not ImportError.
            try:
                import pdfkit

                return pdfkit.from_string(
                    html,
                    False,
                    options={
                        "page-size": "A4",
                        "margin-top": "18mm",
                        "margin-right": "16mm",
                        "margin-bottom": "18mm",
                        "margin-left": "16mm",
                        "encoding": "UTF-8",
                        "enable-local-file-access": None,
                    },
                )
            except Exception as exc:  # noqa: BLE001
                raise BriefingReportError(
                    "PDF rendering is unavailable on this deployment (neither "
                    "WeasyPrint's native libraries nor wkhtmltopdf are present). "
                    "Word and Excel export still work."
                ) from exc

    @staticmethod
    def render_html(report: Dict[str, Any]) -> str:
        """The same document as a page — a shareable link, and what the PDF renders."""
        from flask import render_template

        return render_template("exports/ea_briefing_report.html", **report, for_pdf=False)

    @classmethod
    def render_docx(cls, report: Dict[str, Any]) -> bytes:
        """Word, because a briefing gets annotated before it gets circulated."""
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - python-docx is pinned
            raise BriefingReportError(
                "Word export is unavailable on this deployment (python-docx is missing)."
            ) from exc

        briefing = report["briefing"]
        document = Document()

        document.add_heading("Enterprise Architecture Briefing", level=0)
        subtitle = document.add_paragraph()
        subtitle.add_run(report.get("organisation") or "Enterprise architecture").bold = True
        subtitle.add_run(
            f"  ·  Briefed {_date(briefing['briefed_at'])}"
            f"  ·  Exported {report['generated_at'].strftime('%d %B %Y')}"
        )

        if briefing["headline"]:
            document.add_heading(briefing["headline"], level=1)
        if briefing["summary"]:
            document.add_paragraph(briefing["summary"])

        document.add_heading("At a glance", level=1)
        glance = document.add_table(rows=0, cols=2)
        glance.style = "Light Grid Accent 1"
        for label, value in (
            ("Findings", cls._fmt(briefing["finding_count"])),
            ("Needing attention", cls._fmt(briefing["flagged_count"])),
            ("Areas covered", cls._fmt(report["totals"]["categories"])),
            ("Source", (briefing["source"] or EM_DASH)),
        ):
            cells = glance.add_row().cells
            cells[0].text = label
            cells[1].text = value

        if not report["findings"]:
            empty = document.add_paragraph()
            empty.add_run(
                "This briefing recorded no findings. That is the platform "
                "reporting nothing notable, not a failure to look."
            ).italic = True
        else:
            document.add_heading("Findings", level=1)
            for finding in report["findings"]:
                document.add_heading(finding["title"] or "Untitled finding", level=2)
                meta = document.add_paragraph()
                meta.add_run(
                    "  ·  ".join(
                        [
                            f"severity: {finding['severity'] or EM_DASH}",
                            f"area: {finding['category'] or EM_DASH}",
                            f"evidence: {finding['evidence'] or EM_DASH}",
                        ]
                    )
                ).italic = True
                if finding["detail"]:
                    document.add_paragraph(finding["detail"])
                if finding["action_url"]:
                    action = document.add_paragraph()
                    action.add_run(
                        f"{finding['action_label'] or 'Verify'}: {finding['action_url']}"
                    ).italic = True

        document.add_heading("Areas", level=1)
        areas = document.add_table(rows=1, cols=3)
        areas.style = "Light Grid Accent 1"
        for idx, header in enumerate(("Area", "Findings", "Needing attention")):
            areas.rows[0].cells[idx].text = header
        for category in report["categories"]:
            cells = areas.add_row().cells
            cells[0].text = category["name"]
            cells[1].text = f"{category['count']:,}"
            cells[2].text = f"{category['flagged']:,}"

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @classmethod
    def render_xlsx(cls, report: Dict[str, Any]) -> bytes:
        """Excel, because findings get triaged and assigned before they get closed."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Font
        except ImportError as exc:  # pragma: no cover - openpyxl is pinned
            raise BriefingReportError(
                "Excel export is unavailable on this deployment (openpyxl is missing)."
            ) from exc

        briefing = report["briefing"]
        workbook = Workbook()

        findings = workbook.active
        findings.title = "Findings"
        findings.append(
            ["Severity", "Area", "Finding", "Detail", "Evidence", "Where to verify"]
        )
        for cell in findings[1]:
            cell.font = Font(bold=True)
        for finding in report["findings"]:
            findings.append(
                [
                    finding["severity"] or "",
                    finding["category"] or "",
                    finding["title"] or "",
                    finding["detail"] or "",
                    finding["evidence"] or "",
                    finding["action_url"] or "",
                ]
            )
        _widths(findings, [14, 18, 52, 80, 30, 40])
        findings.freeze_panes = "A2"
        for row in findings.iter_rows(min_row=2, min_col=4, max_col=4):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        summary = workbook.create_sheet("Summary")
        summary.append(["Enterprise Architecture Briefing"])
        summary["A1"].font = Font(bold=True, size=14)
        summary.append([])
        for label, value in (
            ("Exported", report["generated_at"].strftime("%Y-%m-%d %H:%M UTC")),
            ("Briefed", _date(briefing["briefed_at"], blank=True)),
            ("Organisation", report.get("organisation") or ""),
            ("Headline", briefing["headline"] or ""),
            ("Summary", briefing["summary"] or ""),
            ("Source", briefing["source"] or ""),
            # Empty, not zero: a briefing whose counts were never written is
            # not a briefing that found nothing.
            ("Findings", _blank(briefing["finding_count"])),
            ("Needing attention", _blank(briefing["flagged_count"])),
            ("Findings listed here", report["totals"]["findings_listed"]),
            ("Areas covered", report["totals"]["categories"]),
        ):
            summary.append([label, value])
        _widths(summary, [26, 100])

        areas = workbook.create_sheet("Areas")
        areas.append(["Area", "Findings", "Needing attention"])
        for cell in areas[1]:
            cell.font = Font(bold=True)
        for category in report["categories"]:
            areas.append([category["name"], category["count"], category["flagged"]])
        _widths(areas, [28, 16, 22])

        history = workbook.create_sheet("History")
        history.append(
            ["Briefing", "Generated", "Source", "Headline", "Findings", "Needing attention"]
        )
        for cell in history[1]:
            cell.font = Font(bold=True)
        for row in report["history"]:
            history.append(
                [
                    row["id"],
                    _date(row["generated_at"], blank=True),
                    row["source"] or "",
                    row["headline"] or "",
                    _blank(row["finding_count"]),
                    _blank(row["flagged_count"]),
                ]
            )
        _widths(history, [12, 18, 14, 60, 14, 22])

        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _fmt(value: Optional[int]) -> str:
        """A count that was never written is an em dash. It is never 0."""
        return EM_DASH if value is None else f"{value:,}"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _date(value, blank: bool = False) -> str:
    if value is None:
        return "" if blank else EM_DASH
    return value.strftime("%d %B %Y")


def _blank(value):
    """Excel wants an empty cell for a missing number, not a zero."""
    return "" if value is None else value


def _widths(sheet, widths) -> None:
    for idx, width in enumerate(widths, start=1):
        sheet.column_dimensions[sheet.cell(row=1, column=idx).column_letter].width = width
