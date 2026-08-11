"""The capability model as a document — PDF, Word and Excel.

The gap this closes: Archie could export a Solution Architecture Document and a
solution deck, and could not export anything a *business* architect produces.
The capability map offered CSV, JSON and a picture. None of those is a
deliverable — a business architect's output is a document that goes to an
executive committee, and "here is a CSV of my capability model" is not that.

One builder, three renderers. ``build()`` gathers the report once — tree,
maturity, domains, application coverage, gaps and owners — and each renderer
formats the same structure, so the PDF, the Word file and the workbook cannot
disagree about the numbers. That matters more than it looks: three exporters
each doing their own queries is three chances to publish a different figure for
the same capability count in the same meeting.

Nothing here invents a value. A capability with no assessment shows an em dash
in the document and an empty cell in the workbook, never a zero, because a zero
in a maturity column reads as a measured score of zero and an executive acts on
it. See ``CLAUDE.md`` on fabricated data, and the frameworks overview for the
same distinction: ``current_maturity_level`` defaults to 1 and is never NULL,
so ``maturity_assessment_date`` is what says an assessment happened.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

EM_DASH = "—"


class CapabilityReportError(Exception):
    """A renderer could not produce its format on this deployment."""


class CapabilityReportService:
    """Build and render the capability model report."""

    # ------------------------------------------------------------------
    # Gathering
    # ------------------------------------------------------------------

    @staticmethod
    def build(organisation_name: Optional[str] = None) -> Dict[str, Any]:
        """Collect everything the report needs, once.

        Runs inside the request's tenant context: BusinessCapability is a
        TenantMixin model, so the organisation predicate is applied by
        do_orm_execute and must not be written here as well.
        """
        from app.models.business_capabilities import BusinessCapability

        capabilities = BusinessCapability.query.order_by(
            BusinessCapability.level, BusinessCapability.name
        ).all()

        known_ids = {c.id for c in capabilities}
        children_by_parent: Dict[int, List[Any]] = {}
        for cap in capabilities:
            if cap.parent_capability_id:
                children_by_parent.setdefault(cap.parent_capability_id, []).append(cap)

        assessed = [c for c in capabilities if c.maturity_assessment_date is not None]
        avg_current = (
            sum(c.current_maturity_level or 0 for c in assessed) / len(assessed)
            if assessed
            else None
        )
        avg_target = (
            sum(c.target_maturity_level or 0 for c in assessed) / len(assessed)
            if assessed
            else None
        )

        domains: Dict[str, Dict[str, Any]] = {}
        for cap in capabilities:
            key = cap.business_domain or "Unassigned"
            bucket = domains.setdefault(
                key, {"name": key, "count": 0, "assessed": 0, "current_total": 0}
            )
            bucket["count"] += 1
            if cap.maturity_assessment_date is not None:
                bucket["assessed"] += 1
                bucket["current_total"] += cap.current_maturity_level or 0
        for bucket in domains.values():
            bucket["avg_current"] = (
                bucket["current_total"] / bucket["assessed"] if bucket["assessed"] else None
            )

        # A root is a capability nothing else parents — not "level == 1". The
        # tree view learned this the hard way; a report that quietly drops 34 of
        # 529 capabilities is worse than one that does not exist.
        roots = [
            c
            for c in capabilities
            if not c.parent_capability_id or c.parent_capability_id not in known_ids
        ]

        def as_node(cap, ancestors):
            kids = [
                k for k in children_by_parent.get(cap.id, []) if k.id not in ancestors
            ]
            return {
                "id": cap.id,
                "code": cap.code,
                "name": cap.name,
                "description": cap.description,
                "level": cap.level,
                "domain": cap.business_domain,
                "category": cap.category,
                "owner": cap.business_owner,
                "importance": cap.strategic_importance,
                "assessed": cap.maturity_assessment_date is not None,
                "current": cap.current_maturity_level if cap.maturity_assessment_date else None,
                "target": cap.target_maturity_level if cap.maturity_assessment_date else None,
                "gap": cap.maturity_gap if cap.maturity_assessment_date else None,
                "children": [as_node(k, ancestors | {cap.id}) for k in kids],
            }

        tree = [as_node(r, frozenset()) for r in roots]

        flat: List[Dict[str, Any]] = []

        def walk(nodes, depth=0):
            for node in nodes:
                row = dict(node)
                row["depth"] = depth
                row.pop("children", None)
                flat.append(row)
                walk(node["children"], depth + 1)

        walk(tree)

        return {
            "organisation": organisation_name,
            "generated_at": datetime.utcnow(),
            "totals": {
                "capabilities": len(capabilities),
                "roots": len(roots),
                "assessed": len(assessed),
                "unassessed": len(capabilities) - len(assessed),
                "domains": len(domains),
                "avg_current": avg_current,
                "avg_target": avg_target,
            },
            "domains": sorted(domains.values(), key=lambda d: -d["count"]),
            "tree": tree,
            "flat": flat,
        }

    # ------------------------------------------------------------------
    # Rendering
    # ------------------------------------------------------------------

    @staticmethod
    def render_pdf(report: Dict[str, Any]) -> bytes:
        """HTML → PDF, following the pattern already used for the SAD export."""
        from flask import render_template

        html = render_template("exports/capability_model_report.html", **report, for_pdf=True)

        try:
            from weasyprint import HTML

            return HTML(string=html).write_pdf()
        except (ImportError, OSError):
            # OSError as well as ImportError: WeasyPrint is a CFFI binding over
            # native GTK/Pango, and when the wheel is installed without those
            # libraries — normal on Windows and on slim images — the import
            # raises OSError, not ImportError.
            try:
                import pdfkit

                return pdfkit.from_string(
                    html,
                    False,
                    options={
                        "page-size": "A4",
                        "margin-top": "18mm",
                        "margin-right": "16mm",
                        "margin-bottom": "18mm",
                        "margin-left": "16mm",
                        "encoding": "UTF-8",
                        "enable-local-file-access": None,
                    },
                )
            except Exception as exc:  # noqa: BLE001
                raise CapabilityReportError(
                    "PDF rendering is unavailable on this deployment (neither "
                    "WeasyPrint's native libraries nor wkhtmltopdf are present). "
                    "Word and Excel export still work."
                ) from exc

    @staticmethod
    def render_html(report: Dict[str, Any]) -> str:
        """The same document as a page — a shareable link, and what the PDF renders."""
        from flask import render_template

        return render_template(
            "exports/capability_model_report.html", **report, for_pdf=False
        )

    @classmethod
    def render_docx(cls, report: Dict[str, Any]) -> bytes:
        """Word, because a deliverable that cannot be edited is a screenshot."""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:  # pragma: no cover - python-docx is pinned
            raise CapabilityReportError(
                "Word export is unavailable on this deployment (python-docx is missing)."
            ) from exc

        totals = report["totals"]
        document = Document()

        document.add_heading("Capability Model", level=0)
        subtitle = document.add_paragraph()
        subtitle.add_run(report.get("organisation") or "Enterprise architecture").bold = True
        subtitle.add_run(
            f"  ·  {report['generated_at'].strftime('%d %B %Y')}"
        )

        document.add_heading("Summary", level=1)
        summary = document.add_table(rows=0, cols=2)
        summary.style = "Light Grid Accent 1"
        for label, value in (
            ("Capabilities", f"{totals['capabilities']:,}"),
            ("Level-1 capabilities", f"{totals['roots']:,}"),
            ("Domains", f"{totals['domains']:,}"),
            ("Assessed for maturity", f"{totals['assessed']:,}"),
            ("Not yet assessed", f"{totals['unassessed']:,}"),
            ("Average current maturity", cls._fmt(totals["avg_current"])),
            ("Average target maturity", cls._fmt(totals["avg_target"])),
        ):
            row = summary.add_row().cells
            row[0].text = label
            row[1].text = value

        if totals["assessed"] == 0:
            note = document.add_paragraph()
            note.add_run(
                "No capability in this model carries a maturity assessment, so no "
                "average is shown. An unassessed capability is not a capability "
                "scored at level 1."
            ).italic = True

        document.add_heading("Domains", level=1)
        domain_table = document.add_table(rows=1, cols=4)
        domain_table.style = "Light Grid Accent 1"
        for idx, heading in enumerate(("Domain", "Capabilities", "Assessed", "Avg current")):
            domain_table.rows[0].cells[idx].text = heading
        for domain in report["domains"]:
            cells = domain_table.add_row().cells
            cells[0].text = domain["name"]
            cells[1].text = f"{domain['count']:,}"
            cells[2].text = f"{domain['assessed']:,}"
            cells[3].text = cls._fmt(domain["avg_current"])

        document.add_heading("Capability hierarchy", level=1)
        for row in report["flat"]:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(12 * row["depth"])
            label = f"{row['code']} · " if row["code"] else ""
            paragraph.add_run(f"{label}{row['name']}").bold = row["depth"] == 0
            detail = []
            if row["domain"]:
                detail.append(row["domain"])
            if row["owner"]:
                detail.append(f"owner: {row['owner']}")
            detail.append(
                f"maturity {cls._fmt(row['current'])} → {cls._fmt(row['target'])}"
            )
            paragraph.add_run("   " + "  ·  ".join(detail)).italic = True

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @classmethod
    def render_xlsx(cls, report: Dict[str, Any]) -> bytes:
        """Excel, because the model gets worked on before it gets presented."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font
        except ImportError as exc:  # pragma: no cover - openpyxl is pinned
            raise CapabilityReportError(
                "Excel export is unavailable on this deployment (openpyxl is missing)."
            ) from exc

        workbook = Workbook()

        sheet = workbook.active
        sheet.title = "Capabilities"
        headers = [
            "Code", "Capability", "Level", "Depth", "Domain", "Category",
            "Owner", "Strategic importance", "Current maturity",
            "Target maturity", "Gap", "Assessed", "Description",
        ]
        sheet.append(headers)
        for cell in sheet[1]:
            cell.font = Font(bold=True)

        for row in report["flat"]:
            sheet.append([
                row["code"] or "",
                ("    " * row["depth"]) + row["name"],
                row["level"],
                row["depth"] + 1,
                row["domain"] or "",
                row["category"] or "",
                row["owner"] or "",
                row["importance"] or "",
                # Empty, not zero: a 0 in a maturity column reads as a measured
                # score and gets acted on.
                row["current"] if row["current"] is not None else "",
                row["target"] if row["target"] is not None else "",
                row["gap"] if row["gap"] is not None else "",
                "yes" if row["assessed"] else "no",
                row["description"] or "",
            ])

        widths = [14, 46, 7, 7, 20, 18, 20, 20, 16, 16, 8, 10, 60]
        for idx, width in enumerate(widths, start=1):
            sheet.column_dimensions[sheet.cell(row=1, column=idx).column_letter].width = width
        sheet.freeze_panes = "A2"

        summary = workbook.create_sheet("Summary")
        totals = report["totals"]
        summary.append(["Capability model summary"])
        summary["A1"].font = Font(bold=True, size=14)
        summary.append([])
        for label, value in (
            ("Generated", report["generated_at"].strftime("%Y-%m-%d %H:%M UTC")),
            ("Capabilities", totals["capabilities"]),
            ("Level-1 capabilities", totals["roots"]),
            ("Domains", totals["domains"]),
            ("Assessed for maturity", totals["assessed"]),
            ("Not yet assessed", totals["unassessed"]),
            ("Average current maturity",
             totals["avg_current"] if totals["avg_current"] is not None else "not assessed"),
            ("Average target maturity",
             totals["avg_target"] if totals["avg_target"] is not None else "not assessed"),
        ):
            summary.append([label, value])
        summary.column_dimensions["A"].width = 30
        summary.column_dimensions["B"].width = 26

        domains = workbook.create_sheet("Domains")
        domains.append(["Domain", "Capabilities", "Assessed", "Average current maturity"])
        for cell in domains[1]:
            cell.font = Font(bold=True)
        for domain in report["domains"]:
            domains.append([
                domain["name"],
                domain["count"],
                domain["assessed"],
                domain["avg_current"] if domain["avg_current"] is not None else "",
            ])
        domains.column_dimensions["A"].width = 30
        for column in ("B", "C", "D"):
            domains.column_dimensions[column].width = 24

        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _fmt(value: Optional[float]) -> str:
        """A missing measurement is an em dash. It is never 0.0."""
        return EM_DASH if value is None else f"{value:.1f}"
