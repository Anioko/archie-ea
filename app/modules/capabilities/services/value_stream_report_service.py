"""A value stream as a document — PDF, Word, Excel and a shareable page.

The same gap the capability-model report closed, one deliverable further along:
a business architect's value-stream work is a stage flow with the capabilities
that execute each stage, and until now it could only be looked at on screen.
The BIZBOK grid is a working surface, not something that goes to an operating
committee.

One builder, four renderers. ``build()`` gathers the stream, its ordered
stages, the capability mapped to each stage and the coverage totals exactly
once, and each renderer formats that same structure — so the PDF, the Word file
and the workbook cannot disagree about how many stages have no capability
behind them.

Nothing here invents a value. Cycle time, quality, stage durations and the
per-mapping impact percentages are all nullable columns: where one is NULL the
document shows an em dash and the workbook leaves the cell empty, never a zero.
A zero cycle time is a measurement; a missing one is not, and the two must not
look the same in a paper somebody signs off.

Nothing is silently dropped either. ``upsert_mapping_cell`` never checks that
``value_stream_stage_id`` belongs to ``value_stream_id``, so a grid cell can be
written against another stream's stage; and a capability can be deleted with its
mappings left behind. Both are counted and reported rather than filtered away —
the capability-model report learned that rule the hard way when a
root-detection bug quietly lost 34 of 529 rows.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

EM_DASH = "—"


class ValueStreamReportError(Exception):
    """A renderer could not produce its format on this deployment."""


class ValueStreamReportService:
    """Build and render the value-stream report."""

    # ------------------------------------------------------------------
    # Gathering
    # ------------------------------------------------------------------

    @staticmethod
    def build(
        value_stream_id: int, organisation_name: Optional[str] = None
    ) -> Optional[Dict[str, Any]]:
        """Collect everything the report needs, once. None if no such stream.

        Runs inside the request's tenant context: ValueStream, ValueStreamStage
        and CapabilityValueStreamMapping are all TenantMixin models, so the
        organisation predicate is applied by do_orm_execute and must not be
        written here as well. ``UnifiedCapability`` is *not* a TenantMixin
        model, so its query carries no predicate — but the ids fed to it come
        from mappings that were tenant-filtered, so nothing outside the
        organisation can be reached through it.

        The lookup is ``filter(...).first()`` rather than ``.get()`` on purpose.
        ``Query.get()`` is scoped only on an identity-map *miss*; on a hit it
        returns the cached object without emitting SQL, so no tenant filter runs
        (see CLAUDE.md). A report is exactly the kind of read where that matters.
        """
        from app.models.unified_capability import (
            CapabilityValueStreamMapping,
            UnifiedCapability,
            ValueStream,
            ValueStreamStage,
        )

        stream = ValueStream.query.filter(ValueStream.id == value_stream_id).first()
        if stream is None:
            return None

        stages = (
            ValueStreamStage.query.filter(
                ValueStreamStage.value_stream_id == value_stream_id
            )
            .order_by(ValueStreamStage.stage_order, ValueStreamStage.name)
            .all()
        )

        mappings = CapabilityValueStreamMapping.query.filter(
            CapabilityValueStreamMapping.value_stream_id == value_stream_id
        ).all()

        capability_ids = sorted({m.capability_id for m in mappings if m.capability_id})
        capabilities_by_id: Dict[int, Any] = {}
        if capability_ids:
            capabilities_by_id = {
                cap.id: cap
                for cap in UnifiedCapability.query.filter(
                    UnifiedCapability.id.in_(capability_ids)
                ).all()
            }

        stage_ids = {stage.id for stage in stages}
        by_stage: Dict[int, List[Dict[str, Any]]] = {}
        unplaced: List[Dict[str, Any]] = []
        dangling = 0

        for mapping in mappings:
            capability = capabilities_by_id.get(mapping.capability_id)
            if capability is None:
                # The capability row is gone but the mapping survives. Counted,
                # not hidden: a coverage figure that quietly excludes it is wrong.
                dangling += 1
                continue
            row = {
                "id": capability.id,
                "code": capability.code,
                "name": capability.name,
                "level": capability.level,
                "support_type": mapping.support_type,
                "support_level": mapping.support_level,
                "contribution": mapping.capability_contribution,
                "impact_level": mapping.impact_level,
                "stage_criticality": mapping.stage_criticality,
                "cycle_time_impact": mapping.cycle_time_impact,
                "quality_impact": mapping.quality_impact,
                "cost_impact": mapping.cost_impact,
                "assessor": mapping.assessor,
                "notes": mapping.assessment_notes,
            }
            if mapping.value_stream_stage_id in stage_ids:
                by_stage.setdefault(mapping.value_stream_stage_id, []).append(row)
            else:
                unplaced.append(row)

        for rows in by_stage.values():
            rows.sort(key=lambda r: (r["name"] or "").lower())
        unplaced.sort(key=lambda r: (r["name"] or "").lower())

        stage_rows: List[Dict[str, Any]] = []
        for stage in stages:
            stage_rows.append(
                {
                    "id": stage.id,
                    "name": stage.name,
                    "description": stage.description,
                    "order": stage.stage_order,
                    "stage_type": stage.stage_type,
                    "customer_facing": bool(stage.customer_facing),
                    "quality_gate": bool(stage.quality_gate),
                    "target_duration": stage.target_duration,
                    "current_duration": stage.current_duration,
                    "duration_variance": _difference(
                        stage.current_duration, stage.target_duration
                    ),
                    "capabilities": by_stage.get(stage.id, []),
                }
            )

        covered_capability_ids = {
            row["id"] for rows in by_stage.values() for row in rows
        }

        return {
            "organisation": organisation_name,
            "generated_at": datetime.utcnow(),
            "stream": {
                "id": stream.id,
                "code": stream.code,
                "name": stream.name,
                "description": stream.description,
                "stream_type": stream.value_stream_type,
                "industry_domain": stream.industry_domain,
                "strategic_importance": stream.strategic_importance,
                "business_owner": stream.business_owner,
                "target_cycle_time": stream.target_cycle_time,
                "current_cycle_time": stream.current_cycle_time,
                "cycle_time_variance": _difference(
                    stream.current_cycle_time, stream.target_cycle_time
                ),
                "quality_target": stream.quality_target,
                "current_quality": stream.current_quality,
                "quality_variance": _difference(
                    stream.current_quality, stream.quality_target
                ),
            },
            "stages": stage_rows,
            "unplaced": unplaced,
            "totals": {
                "stages": len(stage_rows),
                "mappings": len(mappings),
                "capabilities": len(covered_capability_ids),
                "stages_without_capability": sum(
                    1 for row in stage_rows if not row["capabilities"]
                ),
                "unplaced_mappings": len(unplaced),
                "dangling_mappings": dangling,
            },
        }

    # ------------------------------------------------------------------
    # Rendering
    # ------------------------------------------------------------------

    @staticmethod
    def render_pdf(report: Dict[str, Any]) -> bytes:
        """HTML → PDF, following the pattern already used for the SAD export."""
        from flask import render_template

        html = render_template("exports/value_stream_report.html", **report, for_pdf=True)

        try:
            from weasyprint import HTML

            return HTML(string=html).write_pdf()
        except (ImportError, OSError):
            # OSError as well as ImportError: WeasyPrint is a CFFI binding over
            # native GTK/Pango, and when the wheel is installed without those
            # libraries — normal on Windows and on slim images — the import
            # raises OSError, not ImportError.
            try:
                import pdfkit

                return pdfkit.from_string(
                    html,
                    False,
                    options={
                        "page-size": "A4",
                        "margin-top": "18mm",
                        "margin-right": "16mm",
                        "margin-bottom": "18mm",
                        "margin-left": "16mm",
                        "encoding": "UTF-8",
                        "enable-local-file-access": None,
                    },
                )
            except Exception as exc:  # noqa: BLE001
                raise ValueStreamReportError(
                    "PDF rendering is unavailable on this deployment (neither "
                    "WeasyPrint's native libraries nor wkhtmltopdf are present). "
                    "Word and Excel export still work."
                ) from exc

    @staticmethod
    def render_html(report: Dict[str, Any]) -> str:
        """The same document as a page — a shareable link, and what the PDF renders."""
        from flask import render_template

        return render_template(
            "exports/value_stream_report.html", **report, for_pdf=False
        )

    @classmethod
    def render_docx(cls, report: Dict[str, Any]) -> bytes:
        """Word, because a deliverable that cannot be edited is a screenshot."""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:  # pragma: no cover - python-docx is pinned
            raise ValueStreamReportError(
                "Word export is unavailable on this deployment (python-docx is missing)."
            ) from exc

        stream = report["stream"]
        totals = report["totals"]
        document = Document()

        document.add_heading(stream["name"] or "Value Stream", level=0)
        subtitle = document.add_paragraph()
        subtitle.add_run(report.get("organisation") or "Business architecture").bold = True
        subtitle.add_run(f"  ·  Value stream  ·  {report['generated_at'].strftime('%d %B %Y')}")

        if stream["description"]:
            document.add_paragraph(stream["description"])

        document.add_heading("Profile", level=1)
        profile = document.add_table(rows=0, cols=2)
        profile.style = "Light Grid Accent 1"
        for label, value in (
            ("Code", stream["code"] or EM_DASH),
            ("Type", _humanise(stream["stream_type"])),
            ("Industry domain", stream["industry_domain"] or EM_DASH),
            ("Strategic importance", _humanise(stream["strategic_importance"])),
            ("Business owner", stream["business_owner"] or EM_DASH),
            ("Current cycle time (days)", cls._fmt(stream["current_cycle_time"])),
            ("Target cycle time (days)", cls._fmt(stream["target_cycle_time"])),
            ("Cycle time variance (days)", cls._fmt(stream["cycle_time_variance"])),
            ("Current quality (%)", cls._fmt(stream["current_quality"])),
            ("Target quality (%)", cls._fmt(stream["quality_target"])),
        ):
            cells = profile.add_row().cells
            cells[0].text = label
            cells[1].text = value

        document.add_heading("Coverage", level=1)
        coverage = document.add_table(rows=0, cols=2)
        coverage.style = "Light Grid Accent 1"
        for label, value in (
            ("Stages", f"{totals['stages']:,}"),
            ("Distinct capabilities mapped", f"{totals['capabilities']:,}"),
            ("Capability-to-stage mappings", f"{totals['mappings']:,}"),
            ("Stages with no capability", f"{totals['stages_without_capability']:,}"),
        ):
            cells = coverage.add_row().cells
            cells[0].text = label
            cells[1].text = value

        for note in _integrity_notes(report):
            paragraph = document.add_paragraph()
            paragraph.add_run(note).italic = True

        document.add_heading("Stages", level=1)
        if not report["stages"]:
            paragraph = document.add_paragraph()
            paragraph.add_run(
                "This value stream has no stages yet, so there is nothing to map "
                "capabilities to."
            ).italic = True

        for stage in report["stages"]:
            heading = f"{stage['order']}. {stage['name']}" if stage["order"] is not None else stage["name"]
            document.add_heading(heading, level=2)
            if stage["description"]:
                document.add_paragraph(stage["description"])

            detail = document.add_paragraph()
            detail.add_run(
                "  ·  ".join(
                    [
                        f"type: {_humanise(stage['stage_type'])}",
                        f"customer facing: {'yes' if stage['customer_facing'] else 'no'}",
                        f"quality gate: {'yes' if stage['quality_gate'] else 'no'}",
                        f"duration {cls._fmt(stage['current_duration'])} → "
                        f"{cls._fmt(stage['target_duration'])}",
                    ]
                )
            ).italic = True

            if not stage["capabilities"]:
                gap = document.add_paragraph()
                gap.paragraph_format.left_indent = Pt(12)
                gap.add_run(
                    "No capability is mapped to this stage — the stage has no "
                    "assessed execution behind it."
                ).italic = True
                continue

            table = document.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            for idx, header in enumerate(
                ("Capability", "Support", "Level", "Contribution %", "Criticality")
            ):
                table.rows[0].cells[idx].text = header
            for capability in stage["capabilities"]:
                cells = table.add_row().cells
                label = f"{capability['code']} · " if capability["code"] else ""
                cells[0].text = f"{label}{capability['name']}"
                cells[1].text = _humanise(capability["support_type"])
                cells[2].text = cls._fmt(capability["support_level"])
                cells[3].text = cls._fmt(capability["contribution"])
                cells[4].text = _humanise(capability["stage_criticality"])

        if report["unplaced"]:
            document.add_heading("Mapped to a stage that no longer exists", level=1)
            for capability in report["unplaced"]:
                document.add_paragraph(capability["name"], style="List Bullet")

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @classmethod
    def render_xlsx(cls, report: Dict[str, Any]) -> bytes:
        """Excel, because the stream gets worked on before it gets presented."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font
        except ImportError as exc:  # pragma: no cover - openpyxl is pinned
            raise ValueStreamReportError(
                "Excel export is unavailable on this deployment (openpyxl is missing)."
            ) from exc

        stream = report["stream"]
        totals = report["totals"]
        workbook = Workbook()

        stages_sheet = workbook.active
        stages_sheet.title = "Stages"
        stages_sheet.append(
            [
                "Order", "Stage", "Type", "Customer facing", "Quality gate",
                "Current duration", "Target duration", "Duration variance",
                "Capabilities mapped", "Description",
            ]
        )
        for cell in stages_sheet[1]:
            cell.font = Font(bold=True)
        for stage in report["stages"]:
            stages_sheet.append(
                [
                    _blank(stage["order"]),
                    stage["name"],
                    stage["stage_type"] or "",
                    "yes" if stage["customer_facing"] else "no",
                    "yes" if stage["quality_gate"] else "no",
                    # Empty, not zero: a 0 duration is a measurement and a
                    # missing one is not, and they must not look the same.
                    _blank(stage["current_duration"]),
                    _blank(stage["target_duration"]),
                    _blank(stage["duration_variance"]),
                    len(stage["capabilities"]),
                    stage["description"] or "",
                ]
            )
        _widths(stages_sheet, [7, 34, 18, 16, 14, 18, 18, 18, 20, 60])
        stages_sheet.freeze_panes = "A2"

        mapping_sheet = workbook.create_sheet("Capability mapping")
        mapping_sheet.append(
            [
                "Stage order", "Stage", "Capability code", "Capability", "Level",
                "Support type", "Support level", "Contribution %", "Impact level",
                "Stage criticality", "Cycle time impact %", "Quality impact %",
                "Cost impact %", "Assessor", "Notes",
            ]
        )
        for cell in mapping_sheet[1]:
            cell.font = Font(bold=True)

        def _mapping_row(stage_order, stage_name, capability):
            return [
                _blank(stage_order),
                stage_name,
                capability["code"] or "",
                capability["name"],
                _blank(capability["level"]),
                capability["support_type"] or "",
                _blank(capability["support_level"]),
                _blank(capability["contribution"]),
                capability["impact_level"] or "",
                capability["stage_criticality"] or "",
                _blank(capability["cycle_time_impact"]),
                _blank(capability["quality_impact"]),
                _blank(capability["cost_impact"]),
                capability["assessor"] or "",
                capability["notes"] or "",
            ]

        for stage in report["stages"]:
            for capability in stage["capabilities"]:
                mapping_sheet.append(_mapping_row(stage["order"], stage["name"], capability))
        for capability in report["unplaced"]:
            mapping_sheet.append(
                _mapping_row(None, "(stage no longer exists)", capability)
            )
        _widths(mapping_sheet, [12, 30, 16, 40, 8, 16, 14, 16, 14, 18, 20, 18, 16, 20, 50])
        mapping_sheet.freeze_panes = "A2"

        summary = workbook.create_sheet("Summary")
        summary.append([f"{stream['name'] or 'Value stream'} — value stream report"])
        summary["A1"].font = Font(bold=True, size=14)
        summary.append([])
        for label, value in (
            ("Generated", report["generated_at"].strftime("%Y-%m-%d %H:%M UTC")),
            ("Organisation", report.get("organisation") or ""),
            ("Code", stream["code"] or ""),
            ("Type", stream["stream_type"] or ""),
            ("Industry domain", stream["industry_domain"] or ""),
            ("Strategic importance", stream["strategic_importance"] or ""),
            ("Business owner", stream["business_owner"] or ""),
            ("Current cycle time (days)", _blank(stream["current_cycle_time"])),
            ("Target cycle time (days)", _blank(stream["target_cycle_time"])),
            ("Cycle time variance (days)", _blank(stream["cycle_time_variance"])),
            ("Current quality (%)", _blank(stream["current_quality"])),
            ("Target quality (%)", _blank(stream["quality_target"])),
            ("Stages", totals["stages"]),
            ("Distinct capabilities mapped", totals["capabilities"]),
            ("Capability-to-stage mappings", totals["mappings"]),
            ("Stages with no capability", totals["stages_without_capability"]),
            ("Mappings to a deleted stage", totals["unplaced_mappings"]),
            ("Mappings to a deleted capability", totals["dangling_mappings"]),
        ):
            summary.append([label, value])
        _widths(summary, [34, 30])

        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _fmt(value: Optional[float]) -> str:
        """A missing measurement is an em dash. It is never 0."""
        if value is None:
            return EM_DASH
        if isinstance(value, float) and not value.is_integer():
            return f"{value:.1f}"
        return f"{int(value):,}"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _difference(current, target):
    """current - target, or None when either side was never recorded."""
    if current is None or target is None:
        return None
    return current - target


def _humanise(value: Optional[str]) -> str:
    return EM_DASH if not value else str(value).replace("_", " ").capitalize()


def _blank(value):
    """Excel wants an empty cell for a missing number, not a zero."""
    return "" if value is None else value


def _widths(sheet, widths) -> None:
    for idx, width in enumerate(widths, start=1):
        sheet.column_dimensions[sheet.cell(row=1, column=idx).column_letter].width = width


def _integrity_notes(report: Dict[str, Any]) -> List[str]:
    """Statements about what the numbers exclude, where something is off."""
    totals = report["totals"]
    notes = []
    unplaced = totals["unplaced_mappings"]
    if unplaced:
        notes.append(
            f"{unplaced} capability {'mapping points' if unplaced == 1 else 'mappings point'} "
            "at a stage that no longer exists; they are listed separately rather "
            "than counted into any stage."
        )
    dangling = totals["dangling_mappings"]
    if dangling:
        notes.append(
            f"{dangling} {'mapping references' if dangling == 1 else 'mappings reference'} "
            "a capability that has been deleted from the model and could not be named."
        )
    return notes
