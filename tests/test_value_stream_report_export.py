"""A value stream as a deliverable — PDF, Word, Excel, shareable page.

The BIZBOK grid is a working surface. What a business architect hands to an
operating committee is the stage flow with the capability behind each stage,
and — the part that makes it worth reading — the stages with nothing behind
them at all.

The properties these tests exist to protect:

* **The report never invents a measurement.** Cycle time, quality and stage
  duration are nullable columns. A stream nobody has timed must not appear as a
  stream with a cycle time of zero, because a reader cannot tell those apart
  and plans against the number.
* **The report never silently drops a row.** A mapping pointing at a deleted
  stage is reported separately, not filtered out of the coverage figures.
* **One builder feeds all four renderers**, so the PDF, the Word file and the
  workbook cannot disagree about how many stages are uncovered.
"""

from __future__ import annotations

import io
import uuid

import pytest

from app.modules.capabilities.services.value_stream_report_service import (
    ValueStreamReportError,
    ValueStreamReportService,
)


@pytest.fixture
def client(app):
    previous = app.config.get("LOGIN_DISABLED", False)
    app.config["LOGIN_DISABLED"] = True
    try:
        yield app.test_client()
    finally:
        app.config["LOGIN_DISABLED"] = previous


def _stream(db_session, name, **kw):
    from app.models.unified_capability import ValueStream

    stream = ValueStream(name=name, **kw)
    db_session.add(stream)
    db_session.flush()
    return stream


def _stage(db_session, stream, name, order, **kw):
    from app.models.unified_capability import ValueStreamStage

    stage = ValueStreamStage(
        name=name, value_stream_id=stream.id, stage_order=order, **kw
    )
    db_session.add(stage)
    db_session.flush()
    return stage


def _capability(db_session, name, **kw):
    from app.models.unified_capability import UnifiedCapability

    capability = UnifiedCapability(name=name, **kw)
    db_session.add(capability)
    db_session.flush()
    return capability


def _map(db_session, capability, stream, stage_id, **kw):
    from app.models.unified_capability import CapabilityValueStreamMapping

    mapping = CapabilityValueStreamMapping(
        capability_id=capability.id,
        value_stream_id=stream.id,
        value_stream_stage_id=stage_id,
        **kw,
    )
    db_session.add(mapping)
    db_session.flush()
    return mapping


# ---------------------------------------------------------------------------
# What the builder gathers
# ---------------------------------------------------------------------------


def test_an_untimed_value_stream_reports_no_cycle_time_rather_than_zero(
    db_session, make_org, tenant_ctx
):
    org = make_org(f"vs-untimed-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        stream = _stream(db_session, "Order to Cash", code="OTC")
        report = ValueStreamReportService.build(stream.id)

    assert report is not None
    assert report["stream"]["current_cycle_time"] is None, (
        "a stream nobody has timed must not read as a stream that takes zero days"
    )
    assert report["stream"]["target_cycle_time"] is None
    assert report["stream"]["cycle_time_variance"] is None
    assert report["stream"]["current_quality"] is None


def test_a_variance_needs_both_sides_to_have_been_measured(
    db_session, make_org, tenant_ctx
):
    org = make_org(f"vs-variance-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        half = _stream(db_session, "Half measured", current_cycle_time=30)
        both = _stream(db_session, "Both measured", current_cycle_time=30,
                       target_cycle_time=20)
        half_report = ValueStreamReportService.build(half.id)
        both_report = ValueStreamReportService.build(both.id)

    assert half_report["stream"]["cycle_time_variance"] is None, (
        "a variance against a target nobody set is not a variance of 30"
    )
    assert both_report["stream"]["cycle_time_variance"] == 10


def test_a_missing_value_stream_builds_to_none(db_session, make_org, tenant_ctx):
    org = make_org(f"vs-missing-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        assert ValueStreamReportService.build(987654321) is None


def test_capabilities_are_grouped_under_the_stage_they_execute(
    db_session, make_org, tenant_ctx
):
    org = make_org(f"vs-grouping-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        stream = _stream(db_session, "Idea to Product")
        capture = _stage(db_session, stream, "Capture demand", 1)
        _stage(db_session, stream, "Fulfil", 2)
        planning = _capability(db_session, "Demand Planning", code="DP")
        _map(db_session, planning, stream, capture.id, support_type="primary",
             support_level=4, capability_contribution=70)
        report = ValueStreamReportService.build(stream.id)

    stages = {stage["name"]: stage for stage in report["stages"]}
    assert [c["name"] for c in stages["Capture demand"]["capabilities"]] == [
        "Demand Planning"
    ]
    assert stages["Fulfil"]["capabilities"] == []
    assert report["totals"]["stages"] == 2
    assert report["totals"]["capabilities"] == 1
    assert report["totals"]["stages_without_capability"] == 1


def test_a_mapping_pointing_outside_this_streams_stages_is_reported_not_dropped(
    db_session, make_org, tenant_ctx
):
    """A report that quietly loses rows is worse than one that does not exist.

    ``upsert_mapping_cell`` never checks that ``value_stream_stage_id`` belongs
    to ``value_stream_id``, so a grid cell can be written against another
    stream's stage. The row is real and is counted; it simply has no stage in
    this document to sit under, and saying so beats filtering it away.
    """
    org = make_org(f"vs-unplaced-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        stream = _stream(db_session, "Procure to Pay")
        elsewhere = _stream(db_session, "Another Stream")
        live = _stage(db_session, stream, "Requisition", 1)
        foreign = _stage(db_session, elsewhere, "Belongs elsewhere", 1)
        capability = _capability(db_session, "Supplier Management", code="SM")
        _map(db_session, capability, stream, live.id)
        _map(db_session, capability, stream, foreign.id)

        report = ValueStreamReportService.build(stream.id)

    assert report["totals"]["unplaced_mappings"] == 1
    assert [row["name"] for row in report["unplaced"]] == ["Supplier Management"]
    assert report["totals"]["mappings"] == 2
    assert [c["name"] for c in report["stages"][0]["capabilities"]] == [
        "Supplier Management"
    ]


def test_stage_order_drives_the_sequence(db_session, make_org, tenant_ctx):
    org = make_org(f"vs-order-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        stream = _stream(db_session, "Concept to Market")
        _stage(db_session, stream, "Third", 3)
        _stage(db_session, stream, "First", 1)
        _stage(db_session, stream, "Second", 2)
        report = ValueStreamReportService.build(stream.id)

    assert [stage["name"] for stage in report["stages"]] == ["First", "Second", "Third"]


# ---------------------------------------------------------------------------
# The renderers
# ---------------------------------------------------------------------------


def test_the_workbook_leaves_an_unmeasured_duration_cell_empty(
    db_session, make_org, tenant_ctx
):
    from openpyxl import load_workbook

    org = make_org(f"vs-xlsx-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        stream = _stream(db_session, "Hire to Retire")
        _stage(db_session, stream, "Recruit", 1)
        payload = ValueStreamReportService.render_xlsx(
            ValueStreamReportService.build(stream.id)
        )

    workbook = load_workbook(io.BytesIO(payload))
    assert {"Stages", "Capability mapping", "Summary"} <= set(workbook.sheetnames)

    sheet = workbook["Stages"]
    headers = [cell.value for cell in sheet[1]]
    column = headers.index("Current duration") + 1
    assert sheet.cell(row=2, column=column).value in (None, ""), (
        "a 0 in a duration column reads as a measured duration of zero"
    )

    summary = {
        row[0]: row[1]
        for row in workbook["Summary"].iter_rows(values_only=True)
        if row[0]
    }
    assert summary["Current cycle time (days)"] in (None, "")


def test_the_word_document_names_the_stage_and_its_capability(
    db_session, make_org, tenant_ctx
):
    from docx import Document

    org = make_org(f"vs-docx-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        stream = _stream(db_session, "Quote to Cash")
        stage = _stage(db_session, stream, "Quotation", 1)
        capability = _capability(db_session, "Pricing", code="PRC")
        _map(db_session, capability, stream, stage.id)
        payload = ValueStreamReportService.render_docx(
            ValueStreamReportService.build(stream.id)
        )

    document = Document(io.BytesIO(payload))
    text = "\n".join(p.text for p in document.paragraphs)
    cells = {cell.text for table in document.tables for row in table.rows for cell in row.cells}
    assert "Quote to Cash" in text
    assert "1. Quotation" in text
    assert "PRC · Pricing" in cells


def test_the_html_report_flags_a_stage_with_no_capability_behind_it(
    app, db_session, make_org, tenant_ctx
):
    org = make_org(f"vs-html-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        stream = _stream(db_session, "Plan to Produce")
        _stage(db_session, stream, "Uncovered stage", 1)
        report = ValueStreamReportService.build(stream.id)
        with app.test_request_context("/"):
            html = ValueStreamReportService.render_html(report)

    assert "Uncovered stage" in html
    assert "No capability is mapped to this stage" in html
    assert "—" in html


# ---------------------------------------------------------------------------
# The route
# ---------------------------------------------------------------------------


@pytest.fixture
def seeded_stream(db_session, make_org, tenant_ctx):
    """A value stream with one covered and one uncovered stage."""
    org = make_org(f"vs-route-{uuid.uuid4().hex[:6]}")
    with tenant_ctx(org.id):
        stream = _stream(db_session, "Route Stream", code="RTE")
        stage = _stage(db_session, stream, "Covered", 1)
        _stage(db_session, stream, "Uncovered", 2)
        capability = _capability(db_session, "Routed Capability", code="RC")
        _map(db_session, capability, stream, stage.id)
    return stream.id


@pytest.mark.parametrize("fmt", ["html", "xlsx", "docx"])
def test_each_format_downloads(client, seeded_stream, fmt):
    resp = client.get(f"/value-streams/{seeded_stream}/report.{fmt}")
    assert resp.status_code == 200, resp.data[:300]
    assert len(resp.data) > 0
    if fmt != "html":
        assert "attachment" in resp.headers.get("Content-Disposition", "")


def test_the_html_response_declares_one_charset(client, seeded_stream):
    resp = client.get(f"/value-streams/{seeded_stream}/report.html")
    assert resp.headers["Content-Type"].count("charset") == 1


def test_an_unsupported_format_is_refused_with_the_list(client, seeded_stream):
    resp = client.get(f"/value-streams/{seeded_stream}/report.rtf")
    assert resp.status_code == 400
    body = resp.get_json()
    assert "rtf" in body["error"]
    assert "xlsx" in body["error"]


def test_a_value_stream_that_does_not_exist_is_a_404(client):
    resp = client.get("/value-streams/987654321/report.html")
    assert resp.status_code == 404
    assert "not found" in resp.get_json()["error"].lower()


def test_a_format_this_deployment_cannot_render_says_so(
    client, seeded_stream, monkeypatch
):
    """WeasyPrint binds native libraries that are absent on Windows and slim images.

    The answer has to name the reason and leave the other three formats
    working — not 500, and not a zero-byte file the user opens and puzzles over.
    """
    from app.modules.capabilities.services import value_stream_report_service as mod

    def _unavailable(_report):
        raise ValueStreamReportError("PDF rendering is unavailable on this deployment.")

    monkeypatch.setattr(
        mod.ValueStreamReportService, "render_pdf", staticmethod(_unavailable)
    )

    resp = client.get(f"/value-streams/{seeded_stream}/report.pdf")
    assert resp.status_code == 503
    assert "unavailable" in resp.get_json()["error"].lower()


def test_the_report_requires_a_login(app, seeded_stream):
    resp = app.test_client().get(f"/value-streams/{seeded_stream}/report.xlsx")
    assert resp.status_code in (302, 401)
